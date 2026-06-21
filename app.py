import io
import re
from datetime import datetime, timezone, timedelta


import streamlit as st
import pandas as pd
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

import Workjournal

IST = timezone(timedelta(hours=5, minutes=30))

st.set_page_config(page_title="WorkLog", layout="wide")
st.title("💻 WorkLog")

# ── AUTH ──────────────────────────────────────────────────────────────────────
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False
if "auth_user" not in st.session_state:
    st.session_state.auth_user = ""

with st.sidebar:
    st.header("👤 Sign In")

    if st.session_state.authenticated:
        st.success(f"Logged in as **{st.session_state.auth_user}**")
        if st.button("🚪 Log Out"):
            st.session_state.authenticated = False
            st.session_state.auth_user = ""
            st.rerun()
    else:
        username_input = st.text_input("Username", placeholder="e.g. aditya").strip().lower()
        pin_input      = st.text_input("PIN", type="password", placeholder="4-digit PIN")

        if username_input and Workjournal.user_exists(username_input):
            # Returning user
            if st.button("Login"):
                if Workjournal.verify_pin(username_input, pin_input):
                    st.session_state.authenticated = True
                    st.session_state.auth_user = username_input
                    st.rerun()
                else:
                    st.error("⛔ Incorrect PIN. Please try again.")
        elif username_input:
            # New user — register
            st.info("New user detected. Confirm your PIN to register.")
            pin_confirm = st.text_input("Confirm PIN", type="password", placeholder="Re-enter PIN")
            if st.button("Create Account"):
                if len(pin_input) < 4:
                    st.error("PIN must be at least 4 digits.")
                elif pin_input != pin_confirm:
                    st.error("⛔ PINs do not match. Please try again.")
                else:
                    Workjournal.set_pin(username_input, pin_input)
                    st.session_state.authenticated = True
                    st.session_state.auth_user = username_input
                    st.rerun()
        else:
            st.warning("Please enter your username to continue.")

if not st.session_state.authenticated:
    st.info("👈 Please sign in from the sidebar to use WorkLog 🗂️.")
    st.stop()

username = st.session_state.auth_user

# ── 1. LOG FORM ──────────────────────────────────────────────────────────────
with st.expander("➕ Log New Activity"):
    with st.form("log_form", clear_on_submit=True):
        col1, col2 = st.columns(2)
        description    = col1.text_input("Activity Description")
        category_pick  = col2.selectbox("Category", Workjournal.VALID_CATEGORIES)

        custom_category = st.text_input(
            "Custom Category — fill this if you selected 'other' above",
            placeholder="e.g. training, client-call..."
        )

        # Common times every 30 min (06:00 – 22:30)
        TIME_LIST = [f"{h:02d}:{m:02d}" for h in range(6, 23) for m in (0, 30)]

        col3, col4 = st.columns(2)
        start_pick = col3.selectbox("Start Time", TIME_LIST)
        end_pick   = col4.selectbox("End Time",   TIME_LIST, index=1)

        col5, col6 = st.columns(2)
        start_custom = col5.text_input("Custom Start (overrides list)", placeholder="e.g. 09:47")
        end_custom   = col6.text_input("Custom End   (overrides list)", placeholder="e.g. 10:23")

        if st.form_submit_button("Log Task"):
            time_pattern = re.compile(r"^\d{2}:\d{2}$")
            category = custom_category.strip() if (category_pick == "other" and custom_category.strip()) else category_pick

            # Custom text overrides the selectbox if filled
            start_t = start_custom.strip() if start_custom.strip() else start_pick
            end_t   = end_custom.strip()   if end_custom.strip()   else end_pick

            if not description.strip():
                st.warning("Please enter an activity description.")
            elif not time_pattern.match(start_t) or not time_pattern.match(end_t):
                st.warning("Custom time must be in HH:MM format (e.g. 09:47).")
            else:
                overlap = Workjournal.check_overlap(
                    Workjournal.load_journal(username), start_t, end_t, username
                )
                if overlap["conflict"] and overlap.get("reason") == "end_before_start":
                    st.error(
                        "⛔ Invalid Time Range: End time must be after start time. "
                        "Please correct the times before submitting."
                    )
                elif overlap["conflict"] and overlap.get("reason") == "overlap":
                    conflict_entry = overlap["with"]
                    c_start = Workjournal.fmt_time(conflict_entry["start_time"])
                    c_end   = Workjournal.fmt_time(conflict_entry["end_time"])
                    c_act   = conflict_entry["activity_description"]
                    st.error(
                        f"⛔ Time Conflict Detected: The entry **{start_t} – {end_t}** overlaps "
                        f"with an existing entry **[{c_start} – {c_end} | {c_act}]**. "
                        f"Overlapping entries are not allowed to ensure accurate time tracking. "
                        f"Please adjust your start or end time."
                    )
                else:
                    Workjournal.log_task(
                        description,
                        category,
                        start_time=start_t,
                        end_time=end_t,
                        username=username,
                    )
                    st.rerun()

# ── 2. BUILD DISPLAY DATA ────────────────────────────────────────────────────
st.header("📋 Activity Log")
entries = Workjournal.load_journal(username)

if entries:
    display_data = []
    last_date = None

    for idx, e in enumerate(entries, 1):
        current_date = Workjournal.fmt_date(e["timestamp"])

        if last_date and current_date != last_date:
            display_data.append({
                "S.No": "---", "Date": "---", "Start": "---",
                "End": "---", "Duration": "---", "Category": "---",
                "Activity": f"--- End of {last_date} / Start of {current_date} ---",
            })

        start_ts = e.get("start_time") or e.get("timestamp")
        end_ts   = e.get("end_time")

        display_data.append({
            "S.No":     idx,
            "Date":     current_date,
            "Start":    Workjournal.fmt_time(start_ts) if start_ts else "N/A",
            "End":      Workjournal.fmt_time(end_ts)   if end_ts   else "N/A",
            "Duration": Workjournal.fmt_duration(start_ts, end_ts),
            "Category": e["category"].capitalize(),
            "Activity": e["activity_description"],
        })
        last_date = current_date

    df = pd.DataFrame(display_data)
    st.table(df)

    # ── 3. TOTAL HOURS WIDGET ─────────────────────────────────────────────────
    total_min = sum(
        Workjournal.duration_minutes(
            e.get("start_time") or e.get("timestamp"),
            e.get("end_time")
        )
        for e in entries
    )
    total_h, total_m = divmod(int(total_min), 60)
    st.metric("⏱️ Total Logged Hours (all entries)", f"{total_h}h {total_m}m")

    # ── 4. PRINT / EXPORT SECTION ─────────────────────────────────────────────
    st.divider()
    st.subheader("📤 Export Report")

    if st.button("🖨️ Select Range & Export"):
        st.session_state.show_modal = True

    if st.session_state.get("show_modal"):
        real_indices = [i for i, row in enumerate(display_data) if row["S.No"] != "---"]

        if real_indices:
            col_a, col_b = st.columns(2)
            start_idx = col_a.number_input("From Serial", 1, len(real_indices), 1)
            end_idx   = col_b.number_input("To Serial", start_idx, len(real_indices), len(real_indices))

            df_numeric = df.copy()
            df_numeric["S.No"] = pd.to_numeric(df["S.No"], errors="coerce")
            start_row  = df_numeric[df_numeric["S.No"] == start_idx].index[0]
            end_row    = df_numeric[df_numeric["S.No"] == end_idx].index[0]
            report_df  = df.iloc[start_row : end_row + 1]

            # Range total hours
            real_rows = report_df[report_df["S.No"] != "---"]
            range_entries = [
                entries[int(row["S.No"]) - 1]
                for _, row in real_rows.iterrows()
                if str(row["S.No"]).isdigit()
            ]
            range_min = sum(
                Workjournal.duration_minutes(
                    e.get("start_time") or e.get("timestamp"),
                    e.get("end_time")
                )
                for e in range_entries
            )
            rh, rm = divmod(int(range_min), 60)
            st.info(f"⏱️ Total hours for selected range: **{rh}h {rm}m**")

            col_h, col_p, col_x = st.columns(3)

            # HTML ─────────────────────────────────────────────────────────────
            with col_h:
                html_content = f"""
                <html><head><style>
                    table {{ border-collapse: collapse; width: 100%; }}
                    th, td {{ border: 1px solid black; padding: 8px; text-align: left; }}
                    h2 {{ font-family: sans-serif; }}
                </style></head>
                <body>
                    <h2>Status Report (Serial {start_idx} to {end_idx}) — Total: {rh}h {rm}m</h2>
                    {report_df.to_html(index=False)}
                </body></html>
                """
                st.download_button(
                    label="⬇️ Download HTML",
                    data=html_content,
                    file_name=f"Report_{start_idx}_to_{end_idx}.html",
                    mime="text/html",
                )

            # PDF ─────────────────────────────────────────────────────────────
            with col_p:
                def build_pdf(dataframe, title):
                    buf = io.BytesIO()
                    doc = SimpleDocTemplate(buf, pagesize=A4,
                                            leftMargin=20, rightMargin=20,
                                            topMargin=40, bottomMargin=40)
                    styles = getSampleStyleSheet()
                    elems  = [Paragraph(title, styles["Title"]), Spacer(1, 12)]

                    rows = [list(dataframe.columns)] + [[str(c) for c in row] for row in dataframe.values.tolist()]
                    tbl  = Table(rows, repeatRows=1)
                    tbl.setStyle(TableStyle([
                        ("BACKGROUND",    (0, 0), (-1, 0), colors.HexColor("#4F81BD")),
                        ("TEXTCOLOR",     (0, 0), (-1, 0), colors.white),
                        ("FONTNAME",      (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("FONTSIZE",      (0, 0), (-1, -1), 7),
                        ("GRID",          (0, 0), (-1, -1), 0.5, colors.grey),
                        ("ROWBACKGROUNDS",(0, 1), (-1, -1), [colors.white, colors.HexColor("#DEEAF1")]),
                        ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
                    ]))
                    elems.append(tbl)
                    elems.append(Spacer(1, 12))
                    elems.append(Paragraph(f"Total logged hours: {rh}h {rm}m", styles["Normal"]))
                    doc.build(elems)
                    buf.seek(0)
                    return buf.read()

                pdf_bytes = build_pdf(report_df, f"Status Report  (Serial {start_idx} – {end_idx})")
                st.download_button(
                    label="⬇️ Download PDF",
                    data=pdf_bytes,
                    file_name=f"Report_{start_idx}_to_{end_idx}.pdf",
                    mime="application/pdf",
                )

            # EXCEL ────────────────────────────────────────────────────────────
            with col_x:
                # Month picker — built from all entries (not just selected range)
                all_real_entries = [e for e in entries]
                all_months = sorted(set(
                    datetime.fromisoformat(e["timestamp"]).strftime("%B %Y")
                    for e in all_real_entries
                ), key=lambda m: datetime.strptime(m, "%B %Y"))

                selected_month = st.selectbox(
                    "📅 Filter Excel by Month",
                    ["All (use selected range above)"] + all_months
                )

                def build_excel(dataframe, src_entries):
                    buf     = io.BytesIO()
                    real_df = dataframe[dataframe["S.No"] != "---"].copy()

                    # Build a minutes column for summaries
                    real_df = real_df.copy()
                    real_df["_mins"] = [
                        Workjournal.duration_minutes(
                            e.get("start_time") or e.get("timestamp"),
                            e.get("end_time")
                        )
                        for e in src_entries
                    ]

                    def mins_to_str(m):
                        h, mn = divmod(int(m), 60)
                        return f"{h}h {mn}m"

                    def autofit(ws):
                        for col in ws.columns:
                            ws.column_dimensions[col[0].column_letter].width = max(
                                len(str(col[0].value or "")),
                                *(len(str(c.value or "")) for c in col[1:]),
                            ) + 4

                    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
                        # Sheet 1 — All entries
                        out = real_df.drop(columns=["_mins"])
                        out.to_excel(writer, sheet_name="All Entries", index=False)
                        autofit(writer.sheets["All Entries"])

                        # Sheet per day
                        for date_val, group in real_df.groupby("Date", sort=False):
                            sname = str(date_val)[:31]
                            g_out = group.drop(columns=["_mins"])
                            g_out.to_excel(writer, sheet_name=sname, index=False)
                            autofit(writer.sheets[sname])

                        # Daily summary
                        daily = (
                            real_df.groupby("Date")["_mins"]
                            .sum()
                            .reset_index()
                            .rename(columns={"Date": "Date", "_mins": "Total Minutes"})
                        )
                        daily["Total Hours"] = daily["Total Minutes"].apply(mins_to_str)
                        daily.drop(columns=["Total Minutes"]).to_excel(
                            writer, sheet_name="Daily Summary", index=False)
                        autofit(writer.sheets["Daily Summary"])

                        # Weekly summary
                        real_df["_date_obj"] = pd.to_datetime(real_df["Date"], format="%d-%m-%Y")
                        real_df["Week"] = real_df["_date_obj"].dt.to_period("W").astype(str)
                        weekly = (
                            real_df.groupby("Week")["_mins"]
                            .sum()
                            .reset_index()
                        )
                        weekly["Total Hours"] = weekly["_mins"].apply(mins_to_str)
                        weekly.drop(columns=["_mins"]).to_excel(
                            writer, sheet_name="Weekly Summary", index=False)
                        autofit(writer.sheets["Weekly Summary"])

                        # Monthly summary
                        real_df["Month"] = real_df["_date_obj"].dt.to_period("M").astype(str)
                        monthly = (
                            real_df.groupby("Month")["_mins"]
                            .sum()
                            .reset_index()
                        )
                        monthly["Total Hours"] = monthly["_mins"].apply(mins_to_str)
                        monthly.drop(columns=["_mins"]).to_excel(
                            writer, sheet_name="Monthly Summary", index=False)
                        autofit(writer.sheets["Monthly Summary"])

                    buf.seek(0)
                    return buf.read()

                # Apply month filter if selected
                if selected_month == "All (use selected range above)":
                    excel_df      = report_df
                    excel_entries = range_entries
                    excel_fname   = f"Report_{start_idx}_to_{end_idx}.xlsx"
                else:
                    month_entries = [
                        e for e in all_real_entries
                        if datetime.fromisoformat(e["timestamp"]).strftime("%B %Y") == selected_month
                    ]
                    month_display = []
                    for idx2, e in enumerate(month_entries, 1):
                        s_ts = e.get("start_time") or e.get("timestamp")
                        e_ts = e.get("end_time")
                        month_display.append({
                            "S.No":     idx2,
                            "Date":     Workjournal.fmt_date(e["timestamp"]),
                            "Start":    Workjournal.fmt_time(s_ts) if s_ts else "N/A",
                            "End":      Workjournal.fmt_time(e_ts) if e_ts else "N/A",
                            "Duration": Workjournal.fmt_duration(s_ts, e_ts),
                            "Category": e["category"].capitalize(),
                            "Activity": e["activity_description"],
                        })
                    excel_df      = pd.DataFrame(month_display)
                    excel_entries = month_entries
                    excel_fname   = f"Report_{selected_month.replace(' ', '_')}.xlsx"

                excel_bytes = build_excel(excel_df, excel_entries)
                st.download_button(
                    label="⬇️ Download Excel",
                    data=excel_bytes,
                    file_name=excel_fname,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )

    # ── 5. WORD REPORT ────────────────────────────────────────────────────────
    st.divider()
    st.subheader("📝 Generate Word Report")

    report_type = st.selectbox(
        "Select report period",
        ["Daily (Today)", "Weekly (Last 7 days)", "Monthly (This month)"]
    )

    if st.button("Generate Word Report (.docx)"):
        today     = datetime.now(IST).date()
        all_real  = [e for e in entries]  # full list

        if report_type == "Daily (Today)":
            filtered = [e for e in all_real if Workjournal.fmt_date(e["timestamp"]) == today.strftime("%d-%m-%Y")]
            period_label = f"Daily Report — {today.strftime('%d %b %Y')}"
        elif report_type == "Weekly (Last 7 days)":
            cutoff   = today - timedelta(days=6)
            filtered = [
                e for e in all_real
                if datetime.fromisoformat(e["timestamp"]).date() >= cutoff
            ]
            period_label = f"Weekly Report — {cutoff.strftime('%d %b')} to {today.strftime('%d %b %Y')}"
        else:
            filtered = [
                e for e in all_real
                if datetime.fromisoformat(e["timestamp"]).date().replace(day=1) == today.replace(day=1)
            ]
            period_label = f"Monthly Report — {today.strftime('%B %Y')}"

        if not filtered:
            st.warning("No entries found for the selected period.")
        else:
            # Build Word doc
            doc = Document()

            # Title
            title_para = doc.add_heading(period_label, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph("")

            # Table
            headers = ["S.No", "Date", "Start", "End", "Duration", "Category", "Activity"]
            tbl = doc.add_table(rows=1, cols=len(headers))
            tbl.style = "Table Grid"

            # Header row
            hdr_cells = tbl.rows[0].cells
            for i, h in enumerate(headers):
                hdr_cells[i].text = h
                run = hdr_cells[i].paragraphs[0].runs[0]
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                hdr_cells[i].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Blue background
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                tc   = hdr_cells[i]._tc
                tcPr = tc.get_or_add_tcPr()
                shd  = OxmlElement("w:shd")
                shd.set(qn("w:val"),   "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"),  "4F81BD")
                tcPr.append(shd)

            # Data rows
            total_mins_doc = 0
            for idx, e in enumerate(filtered, 1):
                start_ts = e.get("start_time") or e.get("timestamp")
                end_ts   = e.get("end_time")
                dur      = Workjournal.fmt_duration(start_ts, end_ts)
                total_mins_doc += Workjournal.duration_minutes(start_ts, end_ts)

                row_cells = tbl.add_row().cells
                row_data  = [
                    str(idx),
                    Workjournal.fmt_date(e["timestamp"]),
                    Workjournal.fmt_time(start_ts) if start_ts else "N/A",
                    Workjournal.fmt_time(end_ts)   if end_ts   else "N/A",
                    dur,
                    e["category"].capitalize(),
                    e["activity_description"],
                ]
                for i, val in enumerate(row_data):
                    row_cells[i].text = val
                    row_cells[i].paragraphs[0].runs[0].font.size = Pt(9)

            # Total hours
            dh, dm = divmod(int(total_mins_doc), 60)
            doc.add_paragraph("")
            total_para = doc.add_paragraph(f"Total Logged Hours: {dh}h {dm}m")
            total_para.runs[0].bold = True
            total_para.runs[0].font.size = Pt(11)

            # Save to buffer
            docx_buf = io.BytesIO()
            doc.save(docx_buf)
            docx_buf.seek(0)

            st.download_button(
                label="⬇️ Download Word Report",
                data=docx_buf.read(),
                file_name=f"WorkLog_{report_type.split()[0]}_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

else:
    st.info("No logs yet.")
